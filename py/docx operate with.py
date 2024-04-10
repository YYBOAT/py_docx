import os
import sys
import pickle
import re
import codecs
import string
from docx import Document

rawpath='E:\\codeproiect\\py1\\报告模板'
testdate1='2024.07.17'
dic_replace={'2023.06':'2024.07','2023.12':'2024.07','2023.07':'2024.07','2022':'2023'} #py字典，格式为key：value
dic_year={'2023':'2024'}
dic_replace_data_after={'直流低电阻测试仪':'2024.07.17','游标卡尺':'2024.06.20','数字式测量投影仪':'2024.06.20',
                        '自然换气老化试验机':'2024.09.06','电子万能试验机':'2024.06.20','平头千分尺':'2024.06.20',
                        '绝缘电阻测试仪':testdate1,'直流电阻测试仪':testdate1,'全自动变比测试仪':testdate1,'功率分析仪':testdate1,
                        '工频耐压全自动控制台':testdate1,'介损测试仪':testdate1,'绝缘油介电强度测试仪':testdate1,'油介损测试仪':testdate1,
                        '雷电冲击试验装置':testdate1}

restr='"(?:[^"])*'

def document_text_replace(root,file_name,dic_):
    "替换文档中的一个或多个特定字符串 root目录 file_name文件名,dic_字典:key为原数据,value为新数据"
    document=Document(root+'\\'+file_name)
    for para in document.paragraphs:    #换文档
       for key,value in dic_.items():
           if key in para.text:
              runstr=''
              for i,run in enumerate(para.runs):
                
                 runstr+=run.text
                 para.runs[i].clear()
               
                 para.runs[0].text= runstr.replace(key,dic_[key])  
          
    for table in document.tables:    #换表格
       for row in table.rows:
          for cell in row.cells:
            for para in cell.paragraphs:
               for key,value in dic_.items():
                  if key in para.text:
                     runstr=''
                     for i,run in enumerate(para.runs):
                
                       runstr+=run.text
                       para.runs[i].clear()
               
                     para.runs[0].text= runstr.replace(key,dic_[key])  

    for sec in document.sections:     #换页眉
       for table in sec.header.tables:
          for row in table.rows:
            for cell in row.cells:
              for para in cell.paragraphs:
                 for key,value in dic_.items():
                    if key in para.text:
                       runstr=''
                       for i,run in enumerate(para.runs):
                
                         runstr+=run.text
                         para.runs[i].clear()
               
                    para.runs[0].text= runstr.replace(key,dic_[key])  
    
    for sec in document.sections:     #换页脚
       for table in sec.footer.tables:
          for row in table.rows:
            for cell in row.cells:
              for para in cell.paragraphs:
                 for key,value in dic_.items():
                    if key in para.text:
                       runstr=''
                       for i,run in enumerate(para.runs):
                
                         runstr+=run.text
                         para.runs[i].clear()
               
                         para.runs[0].text= runstr.replace(key,dic_[key])  
   
    #new_file_name=file_name.replace('docx','doc') #改另存为doc，实际文件还是docx没用
    new_file_name=file_name.replace('CF2023','CF2024')
    document.save(root+'\\'+new_file_name)



def document_text_replace_in_table(root,file_name,dic_,N_offset):
    "替换文档表格中的一个或多个特定字符串,替换找到key之后的某个cell的数据  root目录 file_name文件名,dic_字典:key为原数据,value为新数据,N_offset为单元格偏移量,正数右偏,负数左偏"
    document=Document(root+'\\'+file_name)
   
          
    for table in document.tables:
       for row in table.rows:
          cell_i=0
          for cell in row.cells:
             for para in cell.paragraphs:
                para_i=0
                for key,value in dic_.items():
                   if key in para.text:
                      runstr=''
                      runs_location=row.cells[cell_i+N_offset].paragraphs[para_i].runs #定位到之前一个单元格,这是一个runs对象
                      for i,run in enumerate(runs_location): 
                         runstr+=run.text  #runstr现在没用上
                         runs_location[i].clear()
               
                      runs_location[0].text= value  #全替换
                                             #runstr.replace(runstr,value) 适用于对想改的单元格里的一部分进行替换
                para_i+=1
             cell_i+=1
    document.save(root+'\\'+file_name) #file_name.replace('docx','doc')用来保存成doc文件，但是没用


for root,dirs,files in os.walk('E:\\codeproiect\\py1\\报告模板'):
   for file in files:
       # 找出文件中以.doc结尾并且不以~$开头的文件(~S是为了排除临时文件的)
       if file.endswith('.docx') and not file.startswith( '.$'):
          print(file)
          document_text_replace(root,file,dic_year) 
          #document_text_replace_in_table(root,file,dic_replace_data_after,5)
   



  
   




#   for key,value in dic_replace.items():
               # if key in para.text: 
          #         for run in para.runs:  
             #            run.text=run.text.replace(key,value) 

# while (tmp not in list(dic_replace.keys()) ) and (count<len(runs)):
#                    tmp +=runs[count].text
#                    runs[count].clear()
#                   count +=1
# 

