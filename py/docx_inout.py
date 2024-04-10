import os
import random
import sys
import pickle
import re
import codecs
import string
from docx import Document
from copy import deepcopy
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
rawpath='E:\\codeproiect\\py1\\报告模板'
list1=["数字大气压力表","FZ-2012B","/","±0.1℃\r±0.1%RH\rU=0.5hPa\rk=2","内蒙古自治区计量测试研究院","2024.12.13","","",""] #添加数字大气压力表
def document_operate_with_table(root,file_name,line_num):
    "简易表格操作  root目录 file_name文件名,line_num添加的行数,"
    document=Document(root+'\\'+file_name)
    copy_index=0 #复制行的位置,同时是添加行的上一行
    table,copy_index=find_str_in_table(document,'序号',2)
    new_table_row=build_new_table_rows(table.rows[copy_index],line_num)
    i=0
    while i < line_num-1:
         #table.rows[copy_index]._tr.addnext(new_table_row._element)#将复制的行插入第copy_index行之前
        
        table.rows[copy_index]._tr.addprevious(new_table_row[line_num-i-2]._element)#将复制的行插入第copy_index行之后、用i-2去掉build_new_table_row末尾的元素，最后一个元素老是出问题
        i+=1
    #table.rows[copy_index]._tr.addprevious(new_table_row[0]._element)                   
    document.save(root+'\\'+file_name) #file_name.replace('docx','doc')用来保存成doc文件，但是没用

def document_table_addrow(root,file_name,strlist,newfilename=""):
    "在表格最后加一行  root目录 file_name文件名,str字符集,"
    document=Document(root+'\\'+file_name)
    copy_index=0 #复制行的位置,同时是添加行的上一行
    table,copy_index=find_str_in_table(document,'工频耐压全自动控制台',2) #往后偏移2行来确定目标行
    if copy_index==-1:
      return

    table.add_row()
    rows_len=len(table.rows)-1
    iii=len(table.rows[rows_len-1].cells)
    i=1
    replace_str_in_cells(table.rows[rows_len].cells[0],str(rows_len))  
    while i<iii:
      replace_str_in_cells(table.rows[rows_len].cells[i],strlist[i-1])  
      i+=1

    document.save("E:\\codeproiect\\py_docx\\docx1\\"+file_name) #file_name.replace('docx','doc')用来保存成doc文件，但是没用

def find_str_in_table(document,str,offset=0):  #找到要复制行的位置 如果找不到，返回最后一个表，位置返回-1
   i=0 ; copy_index=0 ; is_not_find1=False ; is_not_find2=False ; is_not_find3=False
   for table in document.tables:
      copy_index=0
      if len(table.rows)==(i+1):
         is_not_find2=True 
      is_breake=0; action=0 ;ii=0; 
      for row in table.rows:
         iii=0
         if len(table.rows)==(ii+1):
             is_not_find2=True
         for cell in row.cells:
            if len(row.cells)==(iii+1):
               is_not_find3=True

            for para in cell.paragraphs:
               if str in para.text:
                  action=1
                  is_breake=1
                  break   
            iii+=1 
         if is_breake==1: break
         copy_index+=1
         ii+=1                                                 #runs_location=row.cells[cell_i+N_offset].paragraphs[para_i].runs #定位到之前一个单元格,这是一个runs对象
      if action==1:
         copy_index+=offset  #找到复制行位置，有时候需要往前或者往后偏移几行来确定目标行
         return table,copy_index
      if(is_not_find1&is_not_find2&is_not_find3):  #找不复制行位置,return
         return table,-1
   i+=1
   return table,-1    #找不复制行位置,return


def build_new_table_rows( table_row_copy,line_num ):  #构建line_num数量新行 可能需要手动加入一些元素
    table_rows=[] #list
    i=0
    while i < line_num:
      
      replace_str_in_cells(table_row_copy.cells[0],str(i+1)) # 替换序号列
      replace_str_in_cells(table_row_copy.cells[1],'HZBXCAWD')  # 替换试品编号列
      replace_str_in_cells(table_row_copy.cells[10],str('{:.2f}'.format(random.uniform(69,71))))  # 替换试验值列 随机生成的一个浮点数，范围在[a, b)之间
      replace_str_in_cells(table_row_copy.cells[14],'2024.03.05')   # 替换有效期列
      print(table_row_copy.cells[0].paragraphs[0].text)
      table_rows.append(deepcopy(table_row_copy))#复制第insert_index行 
      i+=1
    return table_rows

def build_new_table_row( table_row_copy,str_ ):  #构建line_num数量新行 可能需要手动加入一些元素
       
    table_rows=[] #list
    i=0
   
    replace_str_in_cells(table_row_copy.cells[0],str(i+1)) # 替换序号列
    replace_str_in_cells(table_row_copy.cells[1],'HZBXCAWD')  # 替换试品编号列
    replace_str_in_cells(table_row_copy.cells[10],str('{:.2f}'.format(random.uniform(69,71))))  # 替换试验值列 随机生成的一个浮点数，范围在[a, b)之间
    replace_str_in_cells(table_row_copy.cells[14],'2024.03.05')   # 替换有效期列
    print(table_row_copy.cells[0].paragraphs[0].text)
    table_rows.append(deepcopy(table_row_copy))#复制第insert_index行 
      
    return table_rows


def replace_str_in_cells(cell,str):
   if len(cell.paragraphs)>0:
      for para in cell.paragraphs:
         if len(para.runs)>0:
            for i,run in enumerate(para.runs):
               para.runs[i].clear()
               para.runs[0].text= str
         else:para.add_run(str)
   else:
      cell.add_paragraph
      cell.paragraphs[0].add_run(str)
   
   cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER #居中
   cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #居中
      
         
      
      


for root,dirs,files in os.walk('E:\\codeproiect\\py_docx\\报告模板'):
   for file in files:
       # 找出文件中以.docx结尾并且不以~$开头的文件(~S是为了排除临时文件的)
      if file.endswith('.docx') and not file.startswith( '.$'):
         print(file)
         document_table_addrow(root,file,list1) 

# document_table_addrow('E:\\codeproiect\\py1\\temp1','安全1.docx',list1)

#print(str('{:.2f}'.format(random.uniform(69,71))))











